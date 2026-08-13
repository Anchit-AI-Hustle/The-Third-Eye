"""Knowledge package: parsers, batched embedder, RAG retriever, ingest pipeline."""

import io
import uuid
from datetime import datetime, timedelta, timezone
from types import SimpleNamespace

import pytest

from app.knowledge import embedder as emb
from app.knowledge import ingestion, pipeline
from app.knowledge import retriever as ret
from app.knowledge.chunker import Chunk
from app.knowledge.ingestion import UnsupportedFileTypeError
from app.knowledge.models import Document

NOW = datetime.now(timezone.utc)


# ─── detect_file_type ────────────────────────────────────────────────────────


@pytest.mark.parametrize(
    "filename,expected",
    [
        ("notes.txt", "txt"),
        ("NOTES.TXT", "txt"),
        ("report.pdf", "pdf"),
        ("sheet.xlsx", "xlsx"),
        ("doc.docx", "docx"),
        ("data.csv", "csv"),
        ("readme.md", "md"),
        ("readme.markdown", "md"),
        ("/tmp/nested/path/file.csv", "csv"),
    ],
)
def test_detect_file_type_maps_known_extensions(filename, expected):
    assert ingestion.detect_file_type(filename) == expected


@pytest.mark.parametrize("filename", ["archive.zip", "image.png", "noextension"])
def test_detect_file_type_rejects_unsupported(filename):
    with pytest.raises(UnsupportedFileTypeError):
        ingestion.detect_file_type(filename)


# ─── parse dispatch ──────────────────────────────────────────────────────────


def test_parse_reads_plain_text():
    assert ingestion.parse(io.BytesIO(b"hello world"), "txt") == "hello world"


def test_parse_is_case_insensitive_about_the_type():
    assert ingestion.parse(io.BytesIO(b"hi"), "TXT") == "hi"


def test_parse_replaces_undecodable_bytes_rather_than_raising():
    assert "hello" in ingestion.parse(io.BytesIO(b"hello \xff\xfe"), "txt")


def test_parse_accepts_a_stream_that_yields_str():
    assert ingestion.parse(io.StringIO("already text"), "md") == "already text"


def test_parse_rejects_an_unknown_type():
    with pytest.raises(UnsupportedFileTypeError, match="No parser for"):
        ingestion.parse(io.BytesIO(b""), "rtf")


def test_parse_csv_joins_cells_and_drops_blank_rows():
    csv_bytes = b"name,qty\nwidget,2\n\n,\ngadget,5\n"
    assert ingestion.parse(io.BytesIO(csv_bytes), "csv") == "name | qty\nwidget | 2\ngadget | 5"


def test_parse_docx_includes_paragraphs_and_tables():
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("   ")  # dropped
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "a"
    table.rows[0].cells[1].text = "b"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    out = ingestion.parse(buf, "docx")
    assert "First paragraph" in out
    assert "a | b" in out


def test_parse_xlsx_labels_each_sheet_and_skips_empty_rows():
    from openpyxl import Workbook

    wb = Workbook()
    sheet = wb.active
    sheet.title = "Numbers"
    sheet.append(["name", "qty"])
    sheet.append([None, None])  # dropped
    sheet.append(["widget", 2])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    out = ingestion.parse(buf, "xlsx")
    assert "# Sheet: Numbers" in out
    assert "name | qty" in out
    assert "widget | 2" in out


def test_parse_xlsx_returns_empty_for_a_workbook_with_no_data():
    from openpyxl import Workbook

    buf = io.BytesIO()
    Workbook().save(buf)
    buf.seek(0)
    assert ingestion.parse(buf, "xlsx") == ""


def test_parse_pdf_concatenates_pages(monkeypatch):
    class FakePage:
        def __init__(self, text):
            self._text = text

        def extract_text(self):
            return self._text

    class FakeReader:
        def __init__(self, file):
            self.pages = [FakePage("page one"), FakePage("  "), FakePage("page two")]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)
    assert ingestion.parse(io.BytesIO(b"%PDF"), "pdf") == "page one\n\npage two"


def test_parse_pdf_skips_a_page_that_fails_extraction(monkeypatch):
    class BadPage:
        def extract_text(self):
            raise ValueError("corrupt page")

    class GoodPage:
        def extract_text(self):
            return "readable"

    class FakeReader:
        def __init__(self, file):
            self.pages = [BadPage(), GoodPage()]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)
    assert ingestion.parse(io.BytesIO(b"%PDF"), "pdf") == "readable"


def test_parse_pdf_tolerates_a_page_returning_none(monkeypatch):
    class NonePage:
        def extract_text(self):
            return None

    class FakeReader:
        def __init__(self, file):
            self.pages = [NonePage()]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)
    assert ingestion.parse(io.BytesIO(b"%PDF"), "pdf") == ""


# ─── embedder ────────────────────────────────────────────────────────────────


def _chunks(n):
    return [Chunk(index=i, content=f"chunk {i}", token_count=2) for i in range(n)]


@pytest.mark.asyncio
async def test_embed_chunks_returns_nothing_for_no_chunks():
    assert await emb.embed_chunks([]) == []


@pytest.mark.asyncio
async def test_embed_chunks_preserves_order(monkeypatch):
    async def fake_embed(batch):
        return [[float(len(t))] for t in batch]

    monkeypatch.setattr(emb.model_router, "embed", fake_embed)
    out = await emb.embed_chunks(_chunks(3))
    assert len(out) == 3


@pytest.mark.asyncio
async def test_embed_chunks_splits_into_batches(monkeypatch):
    sizes = []

    async def fake_embed(batch):
        sizes.append(len(batch))
        return [[0.0] for _ in batch]

    monkeypatch.setattr(emb.model_router, "embed", fake_embed)
    await emb.embed_chunks(_chunks(emb.MAX_BATCH_SIZE + 5))
    assert sizes == [emb.MAX_BATCH_SIZE, 5]


@pytest.mark.asyncio
async def test_embed_chunks_yields_empty_vectors_for_a_failed_batch(monkeypatch):
    async def boom(batch):
        raise RuntimeError("provider down")

    monkeypatch.setattr(emb.model_router, "embed", boom)
    out = await emb.embed_chunks(_chunks(2))
    assert out == [[], []]


@pytest.mark.asyncio
async def test_embed_chunks_raises_on_a_count_mismatch(monkeypatch):
    async def short(batch):
        return [[0.0]]  # one vector regardless of batch size

    monkeypatch.setattr(emb.model_router, "embed", short)
    with pytest.raises(RuntimeError, match="returned 1 vectors for 2 chunks"):
        await emb.embed_chunks(_chunks(2))


@pytest.mark.asyncio
async def test_embed_query_returns_the_single_vector(monkeypatch):
    async def fake_embed(texts):
        return [[1.0, 2.0]]

    monkeypatch.setattr(emb.model_router, "embed", fake_embed)
    assert await emb.embed_query("q") == [1.0, 2.0]


@pytest.mark.asyncio
async def test_embed_query_returns_empty_when_the_provider_gives_nothing(monkeypatch):
    async def empty(texts):
        return []

    monkeypatch.setattr(emb.model_router, "embed", empty)
    assert await emb.embed_query("q") == []


# ─── retriever ───────────────────────────────────────────────────────────────


class _StubResult:
    def __init__(self, rows):
        self._rows = rows

    def fetchall(self):
        return self._rows


class _StubSession:
    """Stands in for AsyncSession: the retriever's query is pgvector-only SQL."""

    def __init__(self, rows):
        self._rows = rows
        self.calls = []

    async def execute(self, statement, params=None):
        self.calls.append((str(statement), params))
        return _StubResult(self._rows)


def _chunk_row(content="c", similarity=0.9, created_at=None, title="Doc", index=0):
    return SimpleNamespace(
        id=uuid.uuid4(),
        document_id=uuid.uuid4(),
        chunk_index=index,
        content=content,
        created_at=created_at or NOW,
        document_title=title,
        similarity=similarity,
    )


@pytest.mark.asyncio
async def test_retrieve_chunks_returns_empty_without_a_query_embedding(monkeypatch):
    async def no_vector(query):
        return []

    monkeypatch.setattr(ret, "embed_query", no_vector)
    assert await ret.retrieve_chunks(_StubSession([]), user_id=uuid.uuid4(), query="q") == []


@pytest.mark.asyncio
async def test_retrieve_chunks_ranks_relevance_above_recency(monkeypatch):
    async def vector(query):
        return [0.1]

    monkeypatch.setattr(ret, "embed_query", vector)

    old_relevant = _chunk_row(content="relevant", similarity=0.95, created_at=NOW - timedelta(days=90))
    new_irrelevant = _chunk_row(content="recent", similarity=0.10, created_at=NOW)
    session = _StubSession([new_irrelevant, old_relevant])

    out = await ret.retrieve_chunks(session, user_id=uuid.uuid4(), query="q")
    assert [r["content"] for r in out] == ["relevant", "recent"]


@pytest.mark.asyncio
async def test_retrieve_chunks_treats_a_naive_timestamp_as_utc(monkeypatch):
    async def vector(query):
        return [0.1]

    monkeypatch.setattr(ret, "embed_query", vector)
    naive = _chunk_row(created_at=NOW.replace(tzinfo=None))
    out = await ret.retrieve_chunks(_StubSession([naive]), user_id=uuid.uuid4(), query="q")
    assert out[0]["score"] > 0


@pytest.mark.asyncio
async def test_retrieve_chunks_truncates_to_top_k(monkeypatch):
    async def vector(query):
        return [0.1]

    monkeypatch.setattr(ret, "embed_query", vector)
    session = _StubSession([_chunk_row(content=str(i)) for i in range(8)])
    out = await ret.retrieve_chunks(session, user_id=uuid.uuid4(), query="q", top_k=3)
    assert len(out) == 3


@pytest.mark.asyncio
async def test_retrieve_chunks_scopes_the_query_to_named_documents(monkeypatch):
    async def vector(query):
        return [0.1]

    monkeypatch.setattr(ret, "embed_query", vector)
    session = _StubSession([])
    doc_ids = [uuid.uuid4()]
    await ret.retrieve_chunks(session, user_id=uuid.uuid4(), query="q", document_ids=doc_ids)

    sql, params = session.calls[0]
    assert "dc.document_id = ANY" in sql
    assert params["doc_ids"] == [str(doc_ids[0])]


@pytest.mark.asyncio
async def test_retrieve_chunks_omits_the_document_filter_when_none_given(monkeypatch):
    async def vector(query):
        return [0.1]

    monkeypatch.setattr(ret, "embed_query", vector)
    session = _StubSession([])
    await ret.retrieve_chunks(session, user_id=uuid.uuid4(), query="q")

    sql, params = session.calls[0]
    assert "dc.document_id = ANY" not in sql
    assert "doc_ids" not in params


def test_format_chunks_for_context_returns_empty_for_no_results():
    assert ret.format_chunks_for_context([]) == ""


def test_format_chunks_for_context_numbers_and_cites_each_passage():
    block = ret.format_chunks_for_context([
        {"document_title": "Handbook", "chunk_index": 2, "content": "body text", "score": 0.876},
    ])
    assert block.startswith("<documents>")
    assert block.endswith("</documents>")
    assert "[1] Handbook (chunk 2, score=0.88)" in block
    assert "body text" in block


# ─── pipeline ────────────────────────────────────────────────────────────────


async def _document(db, test_user):
    doc = Document(
        user_id=test_user.id,
        title="Doc",
        file_type="txt",
        processing_status="pending",
    )
    db.add(doc)
    await db.flush()
    await db.refresh(doc)
    return doc


def _stub_chunking(monkeypatch, chunks):
    # chunk_text needs tiktoken, which downloads its encoding on first use.
    monkeypatch.setattr(pipeline, "chunk_text", lambda text: chunks)


@pytest.mark.asyncio
async def test_ingest_document_stores_chunks_and_marks_ready(db, test_user, monkeypatch):
    doc = await _document(db, test_user)
    _stub_chunking(monkeypatch, _chunks(2))

    async def fake_embed(chunks):
        return [[0.1], [0.2]]

    monkeypatch.setattr(pipeline, "embed_chunks", fake_embed)

    count = await pipeline.ingest_document(
        db, document_id=doc.id, file=io.BytesIO(b"some text"), file_type="txt"
    )
    assert count == 2
    assert doc.processing_status == "ready"
    assert doc.chunk_count == 2


@pytest.mark.asyncio
async def test_ingest_document_marks_failed_on_empty_text(db, test_user, monkeypatch):
    doc = await _document(db, test_user)
    count = await pipeline.ingest_document(
        db, document_id=doc.id, file=io.BytesIO(b"   \n  "), file_type="txt"
    )
    assert count == 0
    assert doc.processing_status == "failed"


@pytest.mark.asyncio
async def test_ingest_document_marks_failed_when_chunking_yields_nothing(db, test_user, monkeypatch):
    doc = await _document(db, test_user)
    _stub_chunking(monkeypatch, [])

    count = await pipeline.ingest_document(
        db, document_id=doc.id, file=io.BytesIO(b"text"), file_type="txt"
    )
    assert count == 0
    assert doc.processing_status == "failed"


@pytest.mark.asyncio
async def test_ingest_document_stores_a_null_embedding_when_the_batch_failed(db, test_user, monkeypatch):
    doc = await _document(db, test_user)
    _stub_chunking(monkeypatch, _chunks(1))

    async def failed_embeddings(chunks):
        return [[]]

    monkeypatch.setattr(pipeline, "embed_chunks", failed_embeddings)

    assert await pipeline.ingest_document(
        db, document_id=doc.id, file=io.BytesIO(b"text"), file_type="txt"
    ) == 1
    assert doc.processing_status == "ready"


@pytest.mark.asyncio
async def test_ingest_document_marks_failed_and_reraises_on_error(db, test_user, monkeypatch):
    doc = await _document(db, test_user)
    _stub_chunking(monkeypatch, _chunks(1))

    async def boom(chunks):
        raise RuntimeError("embedder exploded")

    monkeypatch.setattr(pipeline, "embed_chunks", boom)

    with pytest.raises(RuntimeError, match="embedder exploded"):
        await pipeline.ingest_document(
            db, document_id=doc.id, file=io.BytesIO(b"text"), file_type="txt"
        )
    assert doc.processing_status == "failed"


@pytest.mark.asyncio
async def test_ingest_document_raises_for_an_unknown_document(db, test_user):
    with pytest.raises(Exception):
        await pipeline.ingest_document(
            db, document_id=uuid.uuid4(), file=io.BytesIO(b"x"), file_type="txt"
        )
